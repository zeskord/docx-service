# Возвращает поддокумент приложения В.
import io
import base64
from docx.shared import Inches, Mm, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH


def app3_subdoc(data, doc):
    sd = doc.new_subdoc()
    
    picture_group_counter = 0
    for picture_group in data["КартинкиПриложенияВ"]:
        
        picture_group_counter += 1

        table = sd.add_table(rows=0, cols=2)

        counter = 0
        picture_group_length = len(picture_group["ДанныеФайлов"])
        for file_entry in picture_group["ДанныеФайлов"]:
            counter = counter + 1
            is_last_pic = counter == (picture_group_length)
            if counter % 2 != 0:
                row_pictures = table.add_row()
                row_pictures.height = Mm(100)
                row_descriptions = table.add_row()
            else:
                row_pictures = table.rows[len(table.rows) - 2]
                row_descriptions = table.rows[len(table.rows) - 1]
            file_guid = file_entry["ИмяФайла"]
            file_extension = file_entry["Расширение"]
            file_base64 = file_entry["ДанныеФайла"]
            image = base64.b64decode(file_base64)
            filename = f'workdir/{file_guid}.{file_extension}'
            with open(filename, "wb") as file:
                file.write(image)

            # Добавляем картинку в файл.
            current_col_id = 1 if counter % 2 == 0 else 0
            cell_pictures = row_pictures.cells[current_col_id]
            # Если это последняя картинка и она нечетная, то нужно объединить последние ячейки.
            if is_last_pic and current_col_id == 0:
                cell_pictures.merge(row_pictures.cells[current_col_id + 1])
                # print(f'is_last_pic row_pictures: {is_last_pic}')
                # a, b = row_pictures.cells[:2]
                # a.merge(b)
            cell_pictures.vertical_alignment = WD_ALIGN_VERTICAL.BOTTOM
            picture_paaragraph = cell_pictures.paragraphs[0]
            picture_paaragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            picture_paaragraph.paragraph_format.space_after = Pt(0)
            run = picture_paaragraph.add_run()
            picture = run.add_picture(filename, width=Mm(75), height=None)

            cell_descriptions = row_descriptions.cells[current_col_id]
            if is_last_pic and current_col_id == 0:
                # print(f'is_last_pic row_descriptions: {is_last_pic}')
                # a, b = row_descriptions.cells[:2]
                # a.merge(b)
                cell_descriptions.merge(row_descriptions.cells[current_col_id + 1])
            description = file_entry["Описание"]
            par_description = cell_descriptions.paragraphs[0]
            par_description.alignment = WD_ALIGN_PARAGRAPH.CENTER
            description_run = par_description.add_run(description)
            description_run.font.size = Pt(12)
        
        par = sd.add_paragraph(f'Рисунок {picture_group_counter}: {picture_group["ОписаниеРисунка"]}')
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER


    sd.add_paragraph("Конец!!!")
    return sd
